"""
Logica di elaborazione separata dall'interfaccia grafica.
Tutte le funzioni restituiscono (successo: bool, messaggio: str).
"""

from qgis.core import (
    QgsProject,
    QgsSpatialIndex,
    QgsField,
    QgsGeometry,
)
from . import codici_difetti
from .compat import crea_campo, placement_around_point


# ---------------------------------------------------------------------------
# HELPER PARSING – progressive e Comment
# ---------------------------------------------------------------------------

def parse_progressiva_metri(value):
    """
    Converte un valore di progressiva in metri (int).
    Accetta:
      • numeri (già in metri, es. 12300 o 12300.0)
      • stringhe numeriche ("12300", "12300,5")
      • formato chilometrico "12+300" (= 12 km + 300 m → 12300)
    Restituisce None se il valore non è interpretabile.
    """
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return int(round(value))
    s = str(value).strip()
    if not s or s.lower() in ("null", "none", "nan"):
        return None
    s = s.replace(",", ".")
    if "+" in s:
        parts = s.split("+")
        if len(parts) != 2:
            return None
        try:
            return int(float(parts[0])) * 1000 + int(round(float(parts[1])))
        except ValueError:
            return None
    try:
        return int(round(float(s)))
    except ValueError:
        return None


# Parole accettate al posto di 1 / 2 nel Comment (puntuale = 1, diffuso = 2)
_COMMENT_PAROLE = {
    "puntuale": 1, "p": 1,
    "diffuso": 2, "ricorrente": 2, "d": 2, "r": 2,
}


def normalizza_comment(value):
    """
    Converte il valore Comment in 1 o 2.
    Accetta numeri ("1", "2", "1.0", 2.0) e le parole
    "puntuale" (=1) e "diffuso"/"ricorrente" (=2), case-insensitive.
    Restituisce None se non interpretabile.
    """
    if value is None:
        return None
    s = str(value).strip().lower()
    if not s or s in ("null", "none", "nan"):
        return None
    if s in _COMMENT_PAROLE:
        return _COMMENT_PAROLE[s]
    try:
        return int(float(s.replace(",", ".")))
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# HELPER CRS – sceglie Monte Mario 1 o 2 in base alla longitudine
# ---------------------------------------------------------------------------

def _scegli_crs_metrico(lon):
    """
    Restituisce EPSG:3003 (Monte Mario 1, Italia Ovest) o
    EPSG:3004 (Monte Mario 2, Italia Est) in base alla longitudine.
    Confine convenzionale: 12.5°E.
    """
    return "EPSG:3003" if lon < 12.5 else "EPSG:3004"


def _lon_media_layer(layer, crs_layer=None):
    """
    Calcola la longitudine media delle feature del layer in WGS84.
    Usa le prime 20 feature come campione per velocità.
    """
    from qgis.core import (QgsCoordinateTransform, QgsCoordinateTransformContext,
                            QgsCoordinateReferenceSystem)
    wgs84 = QgsCoordinateReferenceSystem("EPSG:4326")
    crs_src = crs_layer or layer.crs()
    transform_to_wgs = None
    if crs_src.authid() != "EPSG:4326":
        transform_to_wgs = QgsCoordinateTransform(crs_src, wgs84, QgsCoordinateTransformContext())

    lons = []
    for i, feat in enumerate(layer.getFeatures()):
        if i >= 20:
            break
        g = feat.geometry()
        if g is None or g.isEmpty():
            continue
        centroid = g.centroid()
        if transform_to_wgs:
            centroid.transform(transform_to_wgs)
        lons.append(centroid.asPoint().x())

    return sum(lons) / len(lons) if lons else 12.5


def _transform_a_metrico(crs_sorgente, crs_metrico_authid):
    """
    Restituisce un QgsCoordinateTransform da crs_sorgente al CRS metrico scelto.
    Se già coincidono, restituisce None.
    """
    from qgis.core import (QgsCoordinateTransform, QgsCoordinateTransformContext,
                            QgsCoordinateReferenceSystem)
    crs_target = QgsCoordinateReferenceSystem(crs_metrico_authid)
    if crs_sorgente.authid() == crs_metrico_authid:
        return None
    return QgsCoordinateTransform(crs_sorgente, crs_target, QgsCoordinateTransformContext())


def _geom_metrica(geom, transform):
    """Riproietta una geometria se transform non è None, altrimenti la restituisce invariata."""
    if transform is None:
        return geom
    g = QgsGeometry(geom)
    g.transform(transform)
    return g


def _proietta_layer_geometrie(layer, transform):
    """Restituisce lista di geometrie del layer riproiettate (o originali se transform è None)."""
    if transform is None:
        return [feat.geometry() for feat in layer.getFeatures()]
    result = []
    for feat in layer.getFeatures():
        g = QgsGeometry(feat.geometry())
        g.transform(transform)
        result.append(g)
    return result



def assegna_progressive(nome_layer_punti, nome_layer_foto, nome_campo_prog_punti, log_fn, progress_fn=None,
                         usa_svincoli=False, layer_svincoli=None, campo_nome_svincolo=None, soglia_distanza=0):
    """
    Assegna il valore del campo progressiva dal layer punti al layer foto,
    cercando il punto più vicino per ogni feature foto.

    :param nome_layer_punti: nome del layer vettoriale con le progressive
    :param nome_layer_foto:  nome del layer vettoriale delle foto
    :param nome_campo_prog_punti: nome del campo progressiva nel layer punti
    :param log_fn: callable(msg: str) per il log nella UI
    :return: (bool, str) successo e messaggio finale
    """
    punti_layers = QgsProject.instance().mapLayersByName(nome_layer_punti)
    foto_layers  = QgsProject.instance().mapLayersByName(nome_layer_foto)

    if not punti_layers:
        return False, f"Layer '{nome_layer_punti}' non trovato nel progetto."
    if not foto_layers:
        return False, f"Layer '{nome_layer_foto}' non trovato nel progetto."

    punti_layer = punti_layers[0]
    foto_layer  = foto_layers[0]

    # ── Scegli CRS metrico in base alla longitudine dei dati ─────────────────
    lon_media = _lon_media_layer(foto_layer)
    crs_metrico_authid = _scegli_crs_metrico(lon_media)
    log_fn(f"  CRS di lavoro: {crs_metrico_authid} (longitudine media: {lon_media:.2f}°)")

    t_foto  = _transform_a_metrico(foto_layer.crs(),  crs_metrico_authid)
    t_punti = _transform_a_metrico(punti_layer.crs(), crs_metrico_authid)

    if t_foto:
        log_fn(f"  ℹ️  Layer foto  ({foto_layer.crs().authid()}) → {crs_metrico_authid}")
    if t_punti:
        log_fn(f"  ℹ️  Layer punti ({punti_layer.crs().authid()}) → {crs_metrico_authid}")
    if not t_foto and not t_punti:
        log_fn(f"  ✔ Entrambi i layer già in {crs_metrico_authid}")

    # ── AGGIUNTA CAMPI MANCANTI ──────────────────────────────────────────────
    campi_nuovi = []
    if "prog" not in [f.name() for f in foto_layer.fields()]:
        campi_nuovi.append(crea_campo("prog", "double"))
    if "svincolo" not in [f.name() for f in foto_layer.fields()]:
        campi_nuovi.append(crea_campo("svincolo", "string"))
    if campi_nuovi:
        log_fn(f"Aggiunta campi: {[c.name() for c in campi_nuovi]}...")
        foto_layer.dataProvider().addAttributes(campi_nuovi)
        foto_layer.updateFields()

    # ── CARICA LAYER SVINCOLI SE RICHIESTO ────────────────────────────────────
    svincoli_layer = None
    if usa_svincoli and layer_svincoli:
        sv_layers = QgsProject.instance().mapLayersByName(layer_svincoli)
        if sv_layers:
            svincoli_layer = sv_layers[0]
            log_fn(f"  ✔ Layer svincoli caricato: {layer_svincoli}")
        else:
            log_fn(f"  ⚠ Layer svincoli '{layer_svincoli}' non trovato, ignorato.")

    # ── INDICE SPAZIALE in CRS metrico ───────────────────────────────────────
    log_fn("Costruzione indice spaziale...")
    from qgis.core import QgsFeature as _QgsFeature, QgsCoordinateReferenceSystem as _QgsCRS
    spatial_index = QgsSpatialIndex()
    punti_metrici = {}  # fid → (geom_metrica, valore_progressiva)
    for feat in punti_layer.getFeatures():
        g = _geom_metrica(feat.geometry(), t_punti)
        f2 = _QgsFeature(feat.id())
        f2.setGeometry(g)
        spatial_index.addFeature(f2)
        punti_metrici[feat.id()] = (g, feat[nome_campo_prog_punti])

    foto_layer.startEditing()
    contatore = 0
    errori = 0
    features = list(foto_layer.getFeatures())
    totale = len(features)

    idx_prog     = foto_layer.fields().indexFromName("prog")
    idx_svincolo = foto_layer.fields().indexFromName("svincolo")

    # transform svincoli → metrico
    t_sv = _transform_a_metrico(svincoli_layer.crs(), crs_metrico_authid) if svincoli_layer else None

    for i, foto in enumerate(features):
        if progress_fn and i % 10 == 0:
            progress_fn(i + 1, totale)

        geom = foto.geometry()
        if geom is None or geom.isEmpty():
            log_fn(f"  ⚠ Feature ID {foto.id()} senza geometria, saltata.")
            errori += 1
            continue

        # Geometria foto in CRS metrico
        geom_calc = _geom_metrica(geom, t_foto)

        sv_attuale = str(foto["svincolo"] or "").strip() if idx_svincolo >= 0 else ""

        if usa_svincoli:
            # Salta se già marcata come svincolo
            if sv_attuale and sv_attuale.lower() != "no":
                log_fn(f"  ↷ Feature {foto.id()}: già marcata come svincolo '{sv_attuale}', prog azzerata.")
                foto_layer.changeAttributeValue(foto.id(), idx_prog, None)
                contatore += 1
                continue

            # Controlla se cade in un poligono svincolo
            if svincoli_layer:
                # geom_calc è già in metrico; i poligoni svincolo vanno riproiettati in metrico
                nome_sv = None
                for sv_feat in svincoli_layer.getFeatures():
                    g_sv = _geom_metrica(sv_feat.geometry(), t_sv)
                    if g_sv and g_sv.contains(geom_calc):
                        nome_sv = str(sv_feat[campo_nome_svincolo] or "").strip() if campo_nome_svincolo else "Svincolo"
                        break
                if nome_sv:
                    foto_layer.changeAttributeValue(foto.id(), idx_svincolo, nome_sv)
                    foto_layer.changeAttributeValue(foto.id(), idx_prog, None)
                    log_fn(f"  ↷ Feature {foto.id()}: cade in svincolo '{nome_sv}', prog azzerata.")
                    contatore += 1
                    continue

        elif soglia_distanza > 0:
            if sv_attuale and sv_attuale.lower() not in ("", "no", "null", "none", "sospetto"):
                log_fn(f"  ↷ Feature {foto.id()}: svincolo '{sv_attuale}', prog azzerata.")
                foto_layer.changeAttributeValue(foto.id(), idx_prog, None)
                contatore += 1
                continue

        # Assegna la progressiva
        nearest_ids = spatial_index.nearestNeighbor(geom_calc.asPoint(), 1)
        if not nearest_ids:
            log_fn(f"  ⚠ Nessun punto trovato per foto ID {foto.id()}.")
            errori += 1
            continue

        geom_punto, progressiva = punti_metrici[nearest_ids[0]]
        # Il campo progressiva del layer punti può essere in metri (12300)
        # o in formato chilometrico ("12+300"): converto sempre in metri.
        if isinstance(progressiva, str):
            parsed = parse_progressiva_metri(progressiva)
            if parsed is not None:
                progressiva = parsed
        foto_layer.changeAttributeValue(foto.id(), idx_prog, progressiva)

        # Soglia distanza SOSPETTO (in metri — CRS metrico garantito)
        if not usa_svincoli and soglia_distanza > 0:
            dist = geom_calc.distance(geom_punto)
            if dist > soglia_distanza:
                foto_layer.changeAttributeValue(foto.id(), idx_svincolo, "SOSPETTO")
                log_fn(f"  ⚠ Feature {foto.id()}: distanza {dist:.0f}m > soglia {soglia_distanza}m → SOSPETTO.")

        contatore += 1

    foto_layer.commitChanges()
    msg = f"Progressive assegnate: {contatore} feature aggiornate."
    if errori:
        msg += f" {errori} feature saltate (vedi log)."
    log_fn(msg)
    return True, msg


# ---------------------------------------------------------------------------
# STEP 4 – Popolamento Capitolo / Sottocapitolo / Descrizione
# ---------------------------------------------------------------------------

def popola_codici(nome_layer_foto, log_fn):
    """
    Legge il campo 'Title' di ogni feature e popola
    Capitolo_1, sottocap_1, descr_1 (e _2 se ci sono 2 codici).

    :param nome_layer_foto: nome del layer foto
    :param log_fn: callable per il log
    :return: (bool, str)
    """
    layers = QgsProject.instance().mapLayersByName(nome_layer_foto)
    if not layers:
        return False, f"Layer '{nome_layer_foto}' non trovato nel progetto."

    layer = layers[0]

    if "Title" not in [f.name() for f in layer.fields()]:
        return False, "Il layer non ha il campo 'Title'."

    # Aggiungi i campi necessari se non esistono
    campi_necessari = [
        ("Capitolo_1", "string"),
        ("sottocap_1", "string"),
        ("descr_1",    "string"),
        ("Capitolo_2", "string"),
        ("sottocap_2", "string"),
        ("descr_2",    "string"),
    ]
    esistenti = [f.name() for f in layer.fields()]
    nuovi = [crea_campo(nome, tipo) for nome, tipo in campi_necessari if nome not in esistenti]
    if nuovi:
        log_fn(f"Aggiunta di {len(nuovi)} campi al layer...")
        layer.dataProvider().addAttributes(nuovi)
        layer.updateFields()

    layer.startEditing()
    aggiornate = 0
    salti = 0

    for feature in layer.getFeatures():
        title = feature["Title"]
        if title is None or str(title).strip() == "" or str(title).lower() == "null":
            log_fn(f"  ⚠ Feature ID {feature.id()}: campo Title vuoto, saltata.")
            salti += 1
            continue

        # Ignora le parole "puntuale"/"diffuso" eventualmente scritte nel Title:
        # indicano il Comment (1/2), non sono codici difetto.
        codici = [c for c in str(title).strip().split()
                  if c.lower() not in _COMMENT_PAROLE]
        if not codici:
            log_fn(f"  ⚠ Feature ID {feature.id()}: Title '{title}' senza codice difetto, saltata.")
            salti += 1
            continue

        for idx, codice in enumerate(codici[:2]):   # max 2 codici
            n = idx + 1
            if codice in codici_difetti.CODICI_DICT:
                cap, sotto, descr = codici_difetti.CODICI_DICT[codice]
                layer.changeAttributeValue(feature.id(), layer.fields().indexFromName(f"Capitolo_{n}"),  cap)
                layer.changeAttributeValue(feature.id(), layer.fields().indexFromName(f"sottocap_{n}"), sotto)
                layer.changeAttributeValue(feature.id(), layer.fields().indexFromName(f"descr_{n}"),    descr)
            else:
                log_fn(f"  ⚠ Feature ID {feature.id()}: codice '{codice}' non trovato nel dizionario.")

        aggiornate += 1

    layer.commitChanges()
    msg = f"Codici popolati: {aggiornate} feature aggiornate."
    if salti:
        msg += f" {salti} feature saltate (Title vuoto)."
    log_fn(msg)
    return True, msg


# ---------------------------------------------------------------------------
# STEP 5 – Generazione report Word
# ---------------------------------------------------------------------------

def genera_report_word(nome_layer_foto, word_path, use_carr, add_list, remove_first_row, log_fn, progress_fn=None, dpi=0):
    """
    Genera il documento Word leggendo i dati direttamente dal layer foto QGIS.

    :param nome_layer_foto:  nome del layer foto nel progetto QGIS
    :param word_path:        percorso file Word template
    :param use_carr:         bool – usa il campo Carreggiata nella descrizione
    :param add_list:         bool – aggiunge elenco foto prima di ogni tabella
    :param remove_first_row: bool – rimuove prima riga di ogni tabella
    :param log_fn:           callable per il log
    :return: (bool, str)
    """
    try:
        from pathlib import Path
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from openpyxl import Workbook
    except ImportError as e:
        return False, (
            f"Libreria mancante: {e}\n"
            "Installa le dipendenze nella console Python di QGIS:\n"
            "  import subprocess, sys\n"
            "  subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', 'openpyxl'])"
        )

    try:
        log_fn(f"📍 Lettura dati dal layer '{nome_layer_foto}'...")
        table_images_map = _carica_immagini_da_layer(nome_layer_foto, use_carr, log_fn, Path)

        if not table_images_map:
            return False, "Nessuna feature valida trovata nel layer. Controlla che Title e Comment siano compilati."

        log_fn("🖼️  Inserimento foto nel documento Word...")
        if dpi and dpi > 0:
            try:
                from PIL import Image as _PILTest
                log_fn(f"🗜  Compressione immagini a {dpi} ppi...")
            except ImportError:
                log_fn("⚠ Pillow non trovata - impossibile comprimere le foto.")
                log_fn("  Installala dalla console Python di QGIS:")
                log_fn("  import subprocess, sys")
                log_fn("  subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'Pillow'])")
                log_fn("  Poi riavvia QGIS e riprova.")
                dpi = 0   # ricade su originale
        out_doc, out_xlsx = _inserisci_foto_in_word(
            word_path, table_images_map, log_fn,
            add_list, remove_first_row,
            Document, Inches, Pt, WD_PARAGRAPH_ALIGNMENT, Workbook,
            progress_fn=progress_fn,
            nome_layer_foto=nome_layer_foto,
            dpi=dpi,
        )
        return True, f"Documenti salvati:\n• {out_doc}\n• {out_xlsx}"

    except Exception as e:
        import traceback
        log_fn(f"❌ {traceback.format_exc()}")
        return False, str(e)


# ---- helpers interni -------------------------------------------------------

def _carica_immagini_da_layer(nome_layer_foto, use_carr, log_fn, Path):
    """Legge le feature dal layer QGIS e costruisce la mappa tabella → immagini."""
    layers = QgsProject.instance().mapLayersByName(nome_layer_foto)
    if not layers:
        raise RuntimeError(f"Layer '{nome_layer_foto}' non trovato nel progetto QGIS.")

    layer = layers[0]
    field_names = [f.name() for f in layer.fields()]

    # Campi obbligatori
    for campo in ("Title", "Comment", "Path", "prog"):
        if campo not in field_names:
            raise RuntimeError(
                f"Il layer '{nome_layer_foto}' non ha il campo '{campo}'.\n"
                f"Assicurati di aver eseguito prima i passi Progressiva e Codici."
            )

    table_images_map = {}

    for feature in layer.getFeatures():
        fid     = str(feature.id())
        title   = feature["Title"]
        comment = feature["Comment"]
        path    = feature["Path"]
        prog    = feature["prog"]

        # Validazione campi sempre obbligatori
        if title is None or str(title).strip() in ("", "NULL"):
            log_fn(f"  ⚠ fid {fid}: Title vuoto, saltato.")
            continue
        if path is None or str(path).strip() in ("", "NULL"):
            log_fn(f"  ⚠ fid {fid}: Path vuoto, saltato.")
            continue
        # Le foto su svincolo hanno prog=NULL per design: non saltarle se hanno svincolo compilato
        sv_check = ""
        if "svincolo" in field_names:
            sv_raw = feature["svincolo"]
            if sv_raw and str(sv_raw).strip().lower() not in ("", "no", "null", "none", "sospetto"):
                sv_check = str(sv_raw).strip()
        if prog is None and not sv_check:
            log_fn(f"  ⚠ fid {fid}: prog vuoto e nessun svincolo (assegna prima le progressive), saltato.")
            continue

        title   = str(title).strip()
        comment = str(comment).strip() if comment is not None else ""

        # Separa i token del Title: i codici difetto e le eventuali parole
        # "puntuale"/"diffuso" (accettate al posto di 1/2 anche se scritte qui).
        comment_da_title = None
        codici_tok = []
        for tok in title.split():
            tl = tok.lower()
            if tl in _COMMENT_PAROLE:
                comment_da_title = _COMMENT_PAROLE[tl]
            else:
                codici_tok.append(tok)

        if not codici_tok:
            log_fn(f"  ⚠ fid {fid}: Title '{title}' non contiene un codice difetto, saltato.")
            continue

        # Estrai solo il PRIMO codice (Title può contenere 2 codici separati da spazio,
        # es. "11.1.2 3.4" oppure "20 1.2") — l'indice tabella si calcola dal primo.
        primo_codice = codici_tok[0]             # es. "11.1.2"
        parti_codice = primo_codice.split(".")   # es. ['11', '1', '2']

        try:
            first_number = int(parti_codice[0])  # es. 11
        except ValueError:
            log_fn(f"  ⚠ fid {fid}: Title '{title}' non inizia con numero, saltato.")
            continue

        # Codice 20 = "Scartare": escludi dal report, Comment non serve
        if first_number == 20:
            log_fn(f"  ↷ fid {fid}: codice 20 (scartare), escluso dal report.")
            continue

        # Calcolo indice tabella Word
        if first_number <= 10:
            # Solo per codici 1-10 il Comment è necessario:
            # 1/"puntuale" oppure 2/"diffuso" (anche "ricorrente")
            comment_value = normalizza_comment(comment)
            if comment_value is None:
                comment_value = comment_da_title
            if comment_value is None:
                log_fn(f"  ⚠ fid {fid}: Comment '{comment}' non valido per codice {primo_codice} "
                       f"(accettati: 1/'puntuale' o 2/'diffuso'), saltato.")
                continue
            if comment_value == 1:
                y = first_number
            elif comment_value == 2:
                y = first_number + 10
            else:
                log_fn(f"  ⚠ fid {fid}: Comment deve essere 1/'puntuale' o 2/'diffuso' (trovato '{comment}'), saltato.")
                continue
        else:
            # Per codici 11.x.y: Comment non viene letto,
            # il secondo numero del codice determina la sottosezione
            if len(parti_codice) > 1:
                try:
                    z = int(parti_codice[1])   # es. 1 da "11.1.2"
                except ValueError:
                    log_fn(f"  ⚠ fid {fid}: secondo valore di '{primo_codice}' non numerico, saltato.")
                    continue
            else:
                log_fn(f"  ⚠ fid {fid}: codice '{primo_codice}' privo di secondo valore, saltato.")
                continue
            y = first_number + 9 + z   # es. 11+9+1 = 21

        # Descrizione
        description = ""
        if "descr_1" in field_names:
            val = feature["descr_1"]
            if val and str(val).strip() not in ("", "NULL"):
                description = str(val).strip()
        if not description:
            description = f"ID: {fid}"

        # Progressiva formattata oppure nome svincolo
        sv_val = ""
        if "svincolo" in field_names:
            sv_raw = feature["svincolo"]
            if sv_raw and str(sv_raw).strip().lower() not in ("", "no", "null", "none", "sospetto"):
                sv_val = str(sv_raw).strip()

        if sv_val:
            # Usa il nome dello svincolo al posto della progressiva
            km_descr = f"Svincolo {sv_val}"
            prog_int = 0
        else:
            # Accetta sia metri (12300) sia formato chilometrico ("12+300")
            prog_int = parse_progressiva_metri(prog)
            if prog_int is None:
                log_fn(f"  ⚠ fid {fid}: progressiva '{prog}' non interpretabile, uso km 0+000.")
                prog_int = 0
            km_part  = prog_int // 1000
            m_part   = prog_int % 1000
            km_descr = f"km {km_part}+{m_part:03d}"

        # Carreggiata
        carr_descr = ""
        if use_carr and "Carr" in field_names:
            carr_val = feature["Carr"]
            if carr_val and str(carr_val).strip() not in ("", "NULL", "None"):
                carr_descr = f"Carr. {str(carr_val).strip()} - "

        desc_completa = f"{carr_descr}{km_descr} - {description}"

        if y not in table_images_map:
            table_images_map[y] = []
        table_images_map[y].append((Path(str(path)), desc_completa, fid, title, comment, prog_int, str(carr_descr)))

    log_fn(f"  → {sum(len(v) for v in table_images_map.values())} foto distribuite in {len(table_images_map)} tabelle.")
    return table_images_map


def _aggiungi_campo_seq(para, font, bookmark_name=None):
    """
    Inserisce un campo SEQ automatico di Word nel paragrafo.
    Se bookmark_name è fornito, avvolge il campo in un segnalibro
    così da poterlo referenziare con un campo REF dall'elenco.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def make_field_run(field_type, text=None):
        r = OxmlElement("w:r")
        # Copia il formato font dal run precedente
        rPr = OxmlElement("w:rPr")
        i_elem = OxmlElement("w:i")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font.name or "Garamond")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int((font.size.pt if font.size else 9) * 2)))
        rPr.append(rFonts)
        rPr.append(i_elem)
        rPr.append(sz)
        r.append(rPr)
        if field_type == "begin":
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "begin")
            r.append(fld)
        elif field_type == "instr":
            instr = OxmlElement("w:instrText")
            instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr.text = text
            r.append(instr)
        elif field_type == "separate":
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "separate")
            r.append(fld)
        elif field_type == "end":
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "end")
            r.append(fld)
        return r

    p = para._p

    if bookmark_name:
        # Apri segnalibro attorno al campo SEQ
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(abs(hash(bookmark_name)) % 100000))
        bm_start.set(qn("w:name"), bookmark_name)
        p.append(bm_start)

    p.append(make_field_run("begin"))
    p.append(make_field_run("instr", " SEQ Foto \\* ARABIC "))
    p.append(make_field_run("separate"))
    r_val = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "0"
    r_val.append(t)
    p.append(r_val)
    p.append(make_field_run("end"))

    if bookmark_name:
        # Chiudi segnalibro
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(abs(hash(bookmark_name)) % 100000))
        p.append(bm_end)


def _aggiungi_campo_ref(para, font, bookmark_name):
    """
    Inserisce un campo REF che punta al segnalibro del campo SEQ
    nella didascalia corrispondente. Con F9 mostra lo stesso numero
    della foto a cui si riferisce.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def make_run(field_type, text=None):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font.name or "Garamond")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int((font.size.pt if font.size else 9) * 2)))
        rPr.append(rFonts)
        rPr.append(sz)
        r.append(rPr)
        if field_type == "begin":
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "begin")
            r.append(fld)
        elif field_type == "instr":
            instr = OxmlElement("w:instrText")
            instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr.text = text
            r.append(instr)
        elif field_type == "separate":
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "separate")
            r.append(fld)
        elif field_type == "end":
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), "end")
            r.append(fld)
        return r

    p = para._p
    p.append(make_run("begin"))
    p.append(make_run("instr", f" REF {bookmark_name} "))
    p.append(make_run("separate"))
    r_val = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "0"
    r_val.append(t)
    p.append(r_val)
    p.append(make_run("end"))


def _comprimi_immagine(image_path, dpi, log_fn=None):
    """
    Ridimensiona l'immagine alla risoluzione target (dpi) usando un buffer
    in memoria — nessun file temporaneo su disco.
    Restituisce un BytesIO oppure il percorso originale se Pillow non è disponibile.
    Larghezza target = 3.15 pollici × dpi (corrisponde alla larghezza di stampa nel Word).
    """
    import io, os
    def _log(msg):
        if log_fn:
            log_fn(msg)
    try:
        from PIL import Image as PILImage
    except ImportError:
        return str(image_path)

    try:
        img = PILImage.open(str(image_path))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        target_w = int(3.15 * dpi)
        w, h = img.size
        if w > target_w:
            ratio = target_w / w
            img = img.resize((target_w, int(h * ratio)), PILImage.LANCZOS)
        buf = io.BytesIO()
        ext = os.path.splitext(str(image_path))[1].lower()
        fmt = "JPEG" if ext in (".jpg", ".jpeg", ".jpe") else "PNG"
        save_kwargs = {"dpi": (dpi, dpi)}
        if fmt == "JPEG":
            save_kwargs["quality"] = 85
            save_kwargs["optimize"] = True
        img.save(buf, format=fmt, **save_kwargs)
        buf.seek(0)
        return buf
    except Exception:
        return str(image_path)


def _inserisci_foto_in_word(word_path, table_images_map, log_fn,
                             add_list, remove_first_row,
                             Document, Inches, Pt, WD_PARAGRAPH_ALIGNMENT, Workbook,
                             progress_fn=None, nome_layer_foto=None, dpi=0):
    doc    = Document(word_path)
    tables = doc.tables

    # Report Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    ws.append(["fid", "Descrizione", "Tabella", "Foto Numero", "Immagine Percorso", "Title", "Comment"])
    for col, w in zip(['A','B','C','D','E','F','G'], [15,50,10,15,50,30,30]):
        ws.column_dimensions[col].width = w

    # Conta totale foto per la barra di progressione
    totale_foto = sum(len(v) for v in table_images_map.values())
    foto_elaborate = 0
    photo_number = 1
    fid_to_num = {}   # accumula {fid (int) → numero_foto} per aggiornare il layer

    # Fix 4: ordina le foto dentro ogni tabella per progressiva crescente
    # La progressiva è codificata nella descrizione (km_descr), ma è più sicuro
    # ordinarla dall'indice originale. Usiamo l'ordine naturale già presente
    # (le feature sono lette dal layer in ordine di fid).
    # Se vuoi ordinare per km, passa prog nella tupla e ordina qui.

    import os as _os
    # Ordina le foto in ogni tabella per progressiva crescente
    for key in table_images_map:
        table_images_map[key].sort(key=lambda x: (x[6], x[5]))

    for table_index, images_info in table_images_map.items():
        if table_index >= len(tables):
            log_fn(f"  ⚠ Tabella {table_index + 1} non esiste nel documento, saltata.")
            continue

        table = tables[table_index]

        # Scarta subito le foto non presenti su disco, così l'elenco e la
        # numerazione restano coerenti con le foto realmente inserite.
        _info_valide = []
        for _info in images_info:
            if _os.path.exists(str(_info[0])):
                _info_valide.append(_info)
            else:
                log_fn(f"  ⚠ Foto non trovata su disco: {_info[0]} (fid: {_info[2]}), saltata.")
        images_info = _info_valide
        if not images_info:
            continue

        # Elenco foto prima della tabella
        if add_list:
            parent = table._element.getparent()
            idx_in_parent = list(parent).index(table._element)
            for idx_list, img_info in enumerate(images_info):
                _, descr, *_ = img_info
                # Crea paragrafo con "• descrizione (Foto " + campo SEQ + ")"
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)
                para.paragraph_format.left_indent  = Inches(0.25)
                # Testo prima del campo
                run_pre = para.add_run(f"• {descr} (Foto ")
                run_pre.font.name = "Garamond"
                run_pre.font.size = Pt(11)
                # Campo REF che punta al segnalibro della didascalia
                _aggiungi_campo_ref(para, run_pre.font, bookmark_name=f"foto_{img_info[2]}")
                # Testo dopo il campo
                run_post = para.add_run(")")
                run_post.font.name = "Garamond"
                run_post.font.size = Pt(11)
                parent.insert(idx_in_parent, para._element)
                idx_in_parent += 1

        if not table.rows:
            log_fn(f"  ⚠ Tabella {table_index + 1} è vuota (nessuna riga), saltata.")
            continue
        if len(table.rows[0].cells) != 2:
            log_fn(f"  ⚠ Tabella {table_index + 1} non ha 2 colonne, saltata.")
            continue

        for i in range(0, len(images_info), 2):
            row_cells = table.add_row().cells
            for j in range(2):
                if i + j >= len(images_info):
                    continue
                image_path, descr, fid, title, comment, _prog, _carr = images_info[i + j]
                # Comprimi se richiesto (restituisce BytesIO o percorso stringa)
                src = _comprimi_immagine(image_path, dpi, log_fn) if dpi and dpi > 0 else str(image_path)
                para = row_cells[j].paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.add_run()
                run.add_picture(src, width=Inches(3.15))

                # Didascalia con campo SEQ automatico di Word
                # Dopo F9 in Word i numeri si aggiornano automaticamente
                desc_run = para.add_run("\nFoto ")
                desc_run.italic = True
                desc_run.font.name = "Garamond"
                desc_run.font.size = Pt(9)

                # Inserisci campo SEQ con segnalibro unico per questa foto
                _aggiungi_campo_seq(para, desc_run.font, bookmark_name=f"foto_{fid}")

                suffix_run = para.add_run(f": {descr}")
                suffix_run.italic = True
                suffix_run.font.name = "Garamond"
                suffix_run.font.size = Pt(9)
                ws.append([fid, descr, table_index, photo_number, str(image_path), title, comment])
                fid_to_num[int(fid)] = photo_number
                foto_elaborate += 1
                if progress_fn:
                    progress_fn(foto_elaborate, totale_foto)
                log_fn(f"  ✔ Foto {photo_number} → tabella {table_index + 1} (fid: {fid})")
                photo_number += 1

        if remove_first_row and table.rows:
            table._tbl.remove(table.rows[0]._tr)

    out_doc  = word_path.replace(".docx", "_updated.docx")
    out_xlsx = word_path.replace(".docx", "_report.xlsx")
    doc.save(out_doc)
    wb.save(out_xlsx)
    log_fn(f"  ✔ Word salvato: {out_doc}")
    log_fn(f"  ✔ Report Excel salvato: {out_xlsx}")

    # Aggiorna il campo foto_num nel layer QGIS
    if nome_layer_foto and fid_to_num:
        layers = QgsProject.instance().mapLayersByName(nome_layer_foto)
        if layers:
            layer = layers[0]
            # Aggiungi campo foto_num se non esiste
            if "foto_num" not in [f.name() for f in layer.fields()]:
                layer.dataProvider().addAttributes([crea_campo("foto_num", "int")])
                layer.updateFields()
            idx = layer.fields().indexFromName("foto_num")
            layer.startEditing()
            for feature in layer.getFeatures():
                num = fid_to_num.get(feature.id())
                layer.changeAttributeValue(feature.id(), idx, num if num is not None else None)
            layer.commitChanges()
            log_fn(f"  ✔ Campo 'foto_num' aggiornato per {len(fid_to_num)} foto.")

    return out_doc, out_xlsx


# ---------------------------------------------------------------------------
# STEP 0 – Importazione dati da Excel nel layer foto
# ---------------------------------------------------------------------------

def importa_dati_da_excel(excel_path, nome_layer_foto,
                           col_nome_foto, col_title, col_comment, col_note, col_carr, col_svincolo,
                           sovrascrivi, log_fn, progress_fn=None):
    """
    Legge un file Excel compilato in campo e copia Title, Comment, Note e Carr
    nel layer foto facendo il match per nome file.

    Strategie di match (in ordine di priorità):
      1. Nome completo senza estensione  (img_0501 == img_0501)
      2. Numero finale con zeri          (img_0501 → 0501 == 0501)
      3. Numero finale senza zeri iniziali (img_0501 → 501 == 501)

    :param excel_path:        percorso file Excel
    :param nome_layer_foto:   nome del layer foto nel progetto
    :param col_nome_foto:     nome colonna Excel con il nome del file foto
    :param col_title:         nome colonna Excel → campo Title
    :param col_comment:       nome colonna Excel → campo Comment
    :param col_note:          nome colonna Excel → campo Note
    :param col_carr:          nome colonna Excel → campo Carr
    :param sovrascrivi:       se False salta le feature con Title già compilato
    :param log_fn:            callable per il log
    :return: (bool, str)
    """
    try:
        import pandas as pd
    except ImportError:
        return False, (
            "Libreria 'pandas' mancante.\n"
            "Installala dalla console Python di QGIS:\n"
            "  import subprocess, sys\n"
            "  subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'pandas', 'openpyxl'])"
        )
    import os

    # ── Carica il layer ──────────────────────────────────────────────────────
    layers = QgsProject.instance().mapLayersByName(nome_layer_foto)
    if not layers:
        return False, f"Layer '{nome_layer_foto}' non trovato nel progetto."
    layer = layers[0]
    field_names = [f.name() for f in layer.fields()]

    if "Path" not in field_names and "name" not in field_names:
        return False, (
            "Il layer foto non ha né il campo 'Path' né 'name'.\n"
            "Assicurati che il layer sia stato importato con il plugin Import Photos."
        )

    # ── Carica Excel ─────────────────────────────────────────────────────────
    try:
        df = pd.read_excel(excel_path)
    except Exception as e:
        return False, f"Errore nella lettura del file Excel: {e}"

    # Normalizza nomi colonne (strip + lowercase per confronto)
    df.columns = df.columns.str.strip()
    col_map = {c.lower(): c for c in df.columns}

    # Solo il nome foto è obbligatorio — tutto il resto è opzionale
    if col_nome_foto.lower() not in col_map:
        return False, (
            f"Colonna nome foto '{col_nome_foto}' non trovata nel file Excel.\n"
            f"Colonne disponibili: {', '.join(df.columns.tolist())}"
        )

    # Tutte le altre colonne sono opzionali: se non trovate si avvisa e si salta
    col_title_r    = col_map.get(col_title.lower()   if col_title.strip()   else "")
    col_comment_r  = col_map.get(col_comment.lower() if col_comment.strip() else "")
    col_note_r     = col_map.get(col_note.lower()    if col_note.strip()    else "")
    col_carr_r     = col_map.get(col_carr.lower()    if col_carr.strip()    else "")
    col_svincolo_r = col_map.get(col_svincolo.lower()if col_svincolo.strip()else "")

    for etichetta, col_orig, col_r in [
        ("codice difetto",       col_title,    col_title_r),
        ("puntuale/ricorrente",  col_comment,  col_comment_r),
        ("note",                 col_note,     col_note_r),
        ("carreggiata",          col_carr,     col_carr_r),
        ("svincolo",             col_svincolo, col_svincolo_r),
    ]:
        if col_orig.strip() and not col_r:
            log_fn(f"  ⚠ Colonna '{col_orig}' ({etichetta}) non trovata nell'Excel — campo non importato.")

    # Usa i nomi reali delle colonne (preservando maiuscole originali)
    col_nome_foto_r = col_map[col_nome_foto.lower()]

    log_fn(f"  {len(df)} righe lette dall'Excel.")

    # ── Costruisce dizionario nome_foto → riga Excel ─────────────────────────
    import re

    def stem(nome):
        """Nome senza estensione, senza spazi, in minuscolo."""
        return os.path.splitext(str(nome).strip())[0].lower()

    def estrai_numero(s):
        """Estrae la sequenza numerica finale di s (con zeri)."""
        m = re.search(r'(\d+)$', s)
        return m.group(1) if m else None

    def normalizza_numero(n):
        """Rimuove gli zeri iniziali per confronto: '0501' → '501'."""
        return str(int(n)) if n else None

    # Tre mappe di lookup:
    #   map_stem    : 'img_0501' → riga
    #   map_num     : '0501'     → riga  (numero con zeri)
    #   map_num_nz  : '501'      → riga  (numero senza zeri iniziali)
    map_stem   = {}
    map_num    = {}
    map_num_nz = {}

    for _, row in df.iterrows():
        s = stem(row[col_nome_foto_r])
        if not s:
            continue
        map_stem[s] = row
        n = estrai_numero(s)
        if n:
            if n not in map_num:
                map_num[n] = row
            nz = normalizza_numero(n)
            if nz and nz not in map_num_nz:
                map_num_nz[nz] = row

    log_fn(f"  {len(map_stem)} nomi foto distinti nell'Excel.")

    # ── Aggiungi campi Note e Carr se non esistono ───────────────────────────
    campi_da_aggiungere = []
    if "Note" not in field_names:
        campi_da_aggiungere.append(crea_campo("Note", "string"))
    if "Carr" not in field_names:
        campi_da_aggiungere.append(crea_campo("Carr", "string"))
    if "svincolo" not in field_names:
        campi_da_aggiungere.append(crea_campo("svincolo", "string"))
    if campi_da_aggiungere:
        log_fn(f"  Aggiunta campi: {[c.name() for c in campi_da_aggiungere]}...")
        layer.dataProvider().addAttributes(campi_da_aggiungere)
        layer.updateFields()
        field_names = [f.name() for f in layer.fields()]

    # ── Match e aggiornamento ────────────────────────────────────────────────
    layer.startEditing()
    aggiornate = 0
    saltate_gia_compilate = 0
    non_trovate = 0
    features_list = list(layer.getFeatures())
    totale_imp = len(features_list)

    for i_imp, feature in enumerate(features_list):
        # Aggiorna barra ogni 10 feature per evitare crash da processEvents durante editing
        if progress_fn and i_imp % 10 == 0:
            progress_fn(i_imp + 1, totale_imp)
        # Ricava il nome file dalla feature (da Path o da name)
        if "Path" in field_names and feature["Path"]:
            nome_file = os.path.basename(str(feature["Path"]))
        elif "name" in field_names and feature["name"]:
            nome_file = str(feature["name"])
        else:
            log_fn(f"  ⚠ Feature ID {feature.id()}: nessun nome file, saltata.")
            non_trovate += 1
            continue

        chiave = stem(nome_file)

        # Match 1: stem completo  (img_0501 == img_0501)
        if chiave in map_stem:
            riga_match = map_stem[chiave]
        else:
            n  = estrai_numero(chiave)       # es. '0501'
            nz = normalizza_numero(n)         # es. '501'
            # Match 2: numero con zeri  (0501 == 0501)
            if n and n in map_num:
                riga_match = map_num[n]
                log_fn(f"  ~ '{nome_file}' abbinato per numero '{n}'.")
            # Match 3: numero senza zeri iniziali  (501 == 501, oppure 0501 matcha 501)
            elif nz and nz in map_num_nz:
                riga_match = map_num_nz[nz]
                log_fn(f"  ~ '{nome_file}' abbinato per numero normalizzato '{nz}'.")
            else:
                log_fn(f"  ⚠ '{nome_file}' non trovato nell'Excel, saltato.")
                non_trovate += 1
                continue

        # Controlla se già compilato (se sovrascrivi=False)
        if not sovrascrivi:
            title_attuale = feature["Title"] if "Title" in field_names else None
            if title_attuale and str(title_attuale).strip() not in ("", "NULL", "None"):
                saltate_gia_compilate += 1
                continue

        riga = riga_match
        # Nota: val() è definita qui dentro il loop ma usa riga locale — ok perché
        # viene chiamata subito nello stesso ciclo e non catturata per dopo
        val = lambda col: ("" if pd.isna(riga[col]) else str(riga[col]).strip())

        if col_title_r and "Title" in field_names:
            layer.changeAttributeValue(
                feature.id(), layer.fields().indexFromName("Title"), val(col_title_r))
        if col_comment_r and "Comment" in field_names:
            # Normalizza "puntuale"/"diffuso" → "1"/"2" (accettati al posto dei numeri)
            comment_raw = val(col_comment_r)
            comment_norm = normalizza_comment(comment_raw)
            layer.changeAttributeValue(
                feature.id(), layer.fields().indexFromName("Comment"),
                str(comment_norm) if comment_norm is not None else comment_raw)
        if col_note_r and "Note" in field_names:
            layer.changeAttributeValue(
                feature.id(), layer.fields().indexFromName("Note"), val(col_note_r))
        if col_svincolo_r and "svincolo" in field_names:
            sv_val = val(col_svincolo_r)
            layer.changeAttributeValue(
                feature.id(), layer.fields().indexFromName("svincolo"), sv_val)
        if col_carr_r and "Carr" in field_names:
            carr_val = val(col_carr_r)
            # Salva il valore così com'è — l'utente sceglie i valori liberi nel tab Progressiva
            layer.changeAttributeValue(
                feature.id(), layer.fields().indexFromName("Carr"), carr_val)

        aggiornate += 1

    layer.commitChanges()

    msg = f"Importazione completata: {aggiornate} foto aggiornate."
    if saltate_gia_compilate:
        msg += f" {saltate_gia_compilate} saltate (già compilate)."
    if non_trovate:
        msg += f" {non_trovate} non trovate nell'Excel (vedi log)."
    log_fn(msg)
    return True, msg


# ---------------------------------------------------------------------------
# STEP 2b – Assegnazione progressiva con carreggiate separate
# ---------------------------------------------------------------------------

def assegna_progressive_carreggiate(nome_layer_foto, campo_carreggiata,
                                     val_carr_a, nome_layer_punti_a, nome_campo_prog_a,
                                     val_carr_b, nome_layer_punti_b, nome_campo_prog_b,
                                     log_fn, progress_fn=None,
                                     usa_svincoli=False, layer_svincoli=None,
                                     campo_nome_svincolo=None, soglia_distanza=0):
    """
    Assegna la progressiva al layer foto usando due layer distinti,
    uno per carreggiata. Legge il campo 'campo_carreggiata' di ogni foto
    e sceglie il layer corrispondente.

    :param nome_layer_foto:     nome layer foto
    :param campo_carreggiata:   nome del campo nel layer foto con il valore carreggiata
    :param val_carr_a:          valore che indica carreggiata A (es. "1", "A", "Nord")
    :param nome_layer_punti_a:  nome layer progressive carreggiata A
    :param val_carr_b:          valore che indica carreggiata B
    :param nome_layer_punti_b:  nome layer progressive carreggiata B
    :param nome_campo_prog:     nome del campo progressiva in entrambi i layer punti
    :param log_fn:              callable per il log
    :return: (bool, str)
    """
    # ── Carica layer foto ─────────────────────────────────────────────────────
    foto_layers = QgsProject.instance().mapLayersByName(nome_layer_foto)
    if not foto_layers:
        return False, f"Layer foto '{nome_layer_foto}' non trovato."
    foto_layer = foto_layers[0]

    if campo_carreggiata not in [f.name() for f in foto_layer.fields()]:
        return False, (
            f"Il campo carreggiata '{campo_carreggiata}' non esiste nel layer foto.\n"
            f"Assicurati di aver compilato il campo prima di procedere."
        )

    # ── Carica layer progressive A e B ───────────────────────────────────────
    layers_a = QgsProject.instance().mapLayersByName(nome_layer_punti_a)
    layers_b = QgsProject.instance().mapLayersByName(nome_layer_punti_b)

    if not layers_a:
        return False, f"Layer progressive A '{nome_layer_punti_a}' non trovato."
    if not layers_b:
        return False, f"Layer progressive B '{nome_layer_punti_b}' non trovato."

    layer_a = layers_a[0]
    layer_b = layers_b[0]

    # ── Controllo campo progressiva ───────────────────────────────────────────
    if nome_campo_prog_a not in [f.name() for f in layer_a.fields()]:
        return False, f"Il campo '{nome_campo_prog_a}' non esiste nel layer progressive A."
    if nome_campo_prog_b not in [f.name() for f in layer_b.fields()]:
        return False, f"Il campo '{nome_campo_prog_b}' non esiste nel layer progressive B."

    # ── Scegli CRS metrico ────────────────────────────────────────────────────
    lon_media = _lon_media_layer(foto_layer)
    crs_metrico_authid = _scegli_crs_metrico(lon_media)
    log_fn(f"  CRS di lavoro: {crs_metrico_authid} (longitudine media: {lon_media:.2f}°)")

    t_foto  = _transform_a_metrico(foto_layer.crs(), crs_metrico_authid)
    t_a     = _transform_a_metrico(layer_a.crs(),    crs_metrico_authid)
    t_b     = _transform_a_metrico(layer_b.crs(),    crs_metrico_authid)

    for nome_l, crs_l in [(nome_layer_punti_a, layer_a.crs()), (nome_layer_punti_b, layer_b.crs())]:
        if crs_l.authid() != crs_metrico_authid:
            log_fn(f"  ℹ️  '{nome_l}' ({crs_l.authid()}) → {crs_metrico_authid}")
    if t_foto:
        log_fn(f"  ℹ️  Layer foto ({foto_layer.crs().authid()}) → {crs_metrico_authid}")

    # ── Indici spaziali in CRS metrico ───────────────────────────────────────
    log_fn("  Costruzione indici spaziali...")
    from qgis.core import QgsFeature as _QgsFeature2

    idx_a = QgsSpatialIndex()
    punti_metrici_a = {}
    for feat in layer_a.getFeatures():
        g = _geom_metrica(feat.geometry(), t_a)
        f2 = _QgsFeature2(feat.id())
        f2.setGeometry(g)
        idx_a.addFeature(f2)
        punti_metrici_a[feat.id()] = (g, feat[nome_campo_prog_a])

    idx_b = QgsSpatialIndex()
    punti_metrici_b = {}
    for feat in layer_b.getFeatures():
        g = _geom_metrica(feat.geometry(), t_b)
        f2 = _QgsFeature2(feat.id())
        f2.setGeometry(g)
        idx_b.addFeature(f2)
        punti_metrici_b[feat.id()] = (g, feat[nome_campo_prog_b])

    # ── Carica layer svincoli se richiesto ────────────────────────────────────
    svincoli_layer = None
    if usa_svincoli and layer_svincoli:
        sv_layers = QgsProject.instance().mapLayersByName(layer_svincoli)
        if sv_layers:
            svincoli_layer = sv_layers[0]
            log_fn(f"  ✔ Layer svincoli caricato: {layer_svincoli}")
        else:
            log_fn(f"  ⚠ Layer svincoli '{layer_svincoli}' non trovato, ignorato.")

    # transform svincoli → metrico (dopo aver caricato il layer)
    t_sv = _transform_a_metrico(svincoli_layer.crs(), crs_metrico_authid) if svincoli_layer else None

    # ── Aggiungi campi mancanti ───────────────────────────────────────────────
    campi_nuovi = []
    if "prog" not in [f.name() for f in foto_layer.fields()]:
        campi_nuovi.append(crea_campo("prog", "double"))
    if "svincolo" not in [f.name() for f in foto_layer.fields()]:
        campi_nuovi.append(crea_campo("svincolo", "string"))
    if campi_nuovi:
        log_fn(f"  Aggiunta campi: {[c.name() for c in campi_nuovi]}...")
        foto_layer.dataProvider().addAttributes(campi_nuovi)
        foto_layer.updateFields()

    # ── Assegnazione ─────────────────────────────────────────────────────────
    foto_layer.startEditing()
    contatore_a = contatore_b = errori = senza_carr = 0
    features = list(foto_layer.getFeatures())
    totale = len(features)
    idx_prog     = foto_layer.fields().indexFromName("prog")
    idx_svincolo = foto_layer.fields().indexFromName("svincolo")

    # Log dei valori attesi per il campo carreggiata
    log_fn(f"  ℹ️  Campo carreggiata: '{campo_carreggiata}' | Valori attesi: '{val_carr_a}' (layer A) / '{val_carr_b}' (layer B)")

    valori_trovati = set()

    for i, foto in enumerate(features):
        if progress_fn and i % 10 == 0:
            progress_fn(i + 1, totale)
        geom = foto.geometry()
        if geom is None or geom.isEmpty():
            log_fn(f"  ⚠ Feature ID {foto.id()} senza geometria, saltata.")
            errori += 1
            continue

        # Geometria foto in CRS metrico
        geom_calc = _geom_metrica(geom, t_foto)

        # Controlla svincolo solo se la gestione svincoli è attiva
        sv_attuale = str(foto["svincolo"] or "").strip() if idx_svincolo >= 0 else ""
        if usa_svincoli or soglia_distanza > 0:
            if sv_attuale and sv_attuale.lower() not in ("", "no", "null", "none", "sospetto"):
                log_fn(f"  ↷ Feature {foto.id()}: svincolo '{sv_attuale}', prog azzerata.")
                foto_layer.changeAttributeValue(foto.id(), idx_prog, None)
                continue

        # Controlla se cade in poligono svincolo (strategia geometrica)
        if usa_svincoli and svincoli_layer:
            # geom_calc è già in metrico; i poligoni svincolo vanno riproiettati
            nome_sv = None
            for sv_feat in svincoli_layer.getFeatures():
                g_sv = _geom_metrica(sv_feat.geometry(), t_sv)
                if g_sv and g_sv.contains(geom_calc):
                    nome_sv = str(sv_feat[campo_nome_svincolo] or "").strip() if campo_nome_svincolo else "Svincolo"
                    break
            if nome_sv:
                foto_layer.changeAttributeValue(foto.id(), idx_svincolo, nome_sv)
                foto_layer.changeAttributeValue(foto.id(), idx_prog, None)
                log_fn(f"  ↷ Feature {foto.id()}: cade in svincolo '{nome_sv}', prog azzerata.")
                continue

        val_carr = str(foto[campo_carreggiata] or "").strip()
        valori_trovati.add(val_carr)
        if not val_carr:
            log_fn(f"  ⚠ Feature ID {foto.id()}: campo carreggiata vuoto, saltata.")
            senza_carr += 1
            continue

        # Scegli indice e dizionario punti metrici in base al valore carreggiata
        if val_carr.lower() == val_carr_a.lower():
            spatial_idx   = idx_a
            punti_m       = punti_metrici_a
        elif val_carr.lower() == val_carr_b.lower():
            spatial_idx   = idx_b
            punti_m       = punti_metrici_b
        else:
            log_fn(f"  ⚠ Feature ID {foto.id()}: valore carreggiata '{val_carr}' non riconosciuto (attesi: '{val_carr_a}' o '{val_carr_b}'), saltata.")
            errori += 1
            continue

        nearest_ids = spatial_idx.nearestNeighbor(geom_calc.asPoint(), 1)
        if not nearest_ids:
            log_fn(f"  ⚠ Nessun punto trovato per foto ID {foto.id()}.")
            errori += 1
            continue

        geom_punto, prog_val = punti_m[nearest_ids[0]]
        if prog_val is None:
            log_fn(f"  ⚠ Feature ID {foto.id()}: campo progressiva del punto più vicino è NULL.")
            errori += 1
            continue

        # Converti in metri se in formato chilometrico ("12+300")
        if isinstance(prog_val, str):
            parsed = parse_progressiva_metri(prog_val)
            if parsed is not None:
                prog_val = parsed

        foto_layer.changeAttributeValue(foto.id(), idx_prog, prog_val)

        # Soglia distanza SOSPETTO (in metri — CRS metrico garantito)
        if soglia_distanza > 0:
            dist = geom_calc.distance(geom_punto)
            if dist > soglia_distanza:
                foto_layer.changeAttributeValue(foto.id(), idx_svincolo, "SOSPETTO")
                log_fn(f"  ⚠ Feature {foto.id()}: distanza {dist:.0f}m > soglia {soglia_distanza}m → SOSPETTO.")

        if val_carr.lower() == val_carr_a.lower():
            contatore_a += 1
        else:
            contatore_b += 1

    foto_layer.commitChanges()

    # Log riassuntivo valori trovati nel campo carreggiata
    if valori_trovati:
        log_fn(f"  ℹ️  Valori trovati nel campo '{campo_carreggiata}': {sorted(valori_trovati)}")

    msg = (f"Progressive assegnate — carr. A: {contatore_a}, carr. B: {contatore_b}.")
    if senza_carr:
        msg += f" {senza_carr} foto senza carreggiata compilata, saltate."
    if errori:
        msg += f" {errori} errori (vedi log)."
    log_fn(msg)
    return True, msg


# ---------------------------------------------------------------------------
# IMPORTA FOTO DA CARTELLA
# ---------------------------------------------------------------------------

def importa_foto_da_cartella(cartella, nome_layer, estensioni, sottocartelle,
                              crs_epsg, data_da, data_a, log_fn, progress_fn=None):
    """
    Scansiona una cartella, legge le coordinate GPS dall'EXIF di ogni foto
    e crea un layer vettoriale puntuale in QGIS.
    """
    import os, datetime, pathlib

    try:
        from PIL import Image as PILImage
        from PIL.ExifTags import TAGS, GPSTAGS
    except ImportError:
        return False, (
            "Pillow non trovata. Installala dalla console Python di QGIS:\n"
            "import subprocess, sys\n"
            "subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'Pillow'])"
        )

    # ── Raccogli tutti i file ────────────────────────────────────────────────
    log_fn(f"  Scansione cartella: {cartella}")
    file_list = []
    if sottocartelle:
        for root, _, files in os.walk(cartella):
            for f in files:
                if os.path.splitext(f)[1].lower() in estensioni:
                    file_list.append(os.path.join(root, f))
    else:
        for f in os.listdir(cartella):
            if os.path.splitext(f)[1].lower() in estensioni:
                file_list.append(os.path.join(cartella, f))

    file_list.sort()
    log_fn(f"  {len(file_list)} file trovati.")

    if not file_list:
        return False, "Nessun file trovato con le estensioni selezionate."

    # ── Helpers EXIF ─────────────────────────────────────────────────────────
    def _get_exif(img):
        try:
            try:
                return img.getexif()   # Pillow moderno: restituisce oggetto Exif
            except AttributeError:
                raw = img._getexif()   # fallback legacy
                if not raw:
                    return {}
                return {TAGS.get(k, k): v for k, v in raw.items()}
        except Exception:
            return {}

    def _get_gps(exif):
        try:
            # Pillow moderno: getexif() restituisce oggetto Exif,
            # GPSInfo è un tag numerico (34853) che punta al sub-IFD
            if hasattr(exif, 'get_ifd'):
                gps_ifd = exif.get_ifd(34853)
                if not gps_ifd:
                    return None
                return {GPSTAGS.get(k, k): v for k, v in gps_ifd.items()}
            # Fallback: exif è un dict normale (da _getexif())
            gps_raw = exif.get("GPSInfo")
            if not gps_raw or not hasattr(gps_raw, 'items'):
                return None
            return {GPSTAGS.get(k, k): v for k, v in gps_raw.items()}
        except Exception:
            return None

    def _dms_to_dd(dms, ref):
        try:
            d = float(dms[0])
            m = float(dms[1])
            s = float(dms[2])
            dd = d + m / 60 + s / 3600
            if ref in ("S", "W"):
                dd = -dd
            return dd
        except Exception:
            return None

    def _parse_datetime(exif):
        # Tag numerici per Pillow moderno (oggetto Exif)
        # 36867=DateTimeOriginal, 306=DateTime, 36868=DateTimeDigitized
        tag_names  = ("DateTimeOriginal", "DateTime", "DateTimeDigitized")
        tag_numbers = (36867, 306, 36868)
        candidates = []
        for tag in tag_names:
            v = exif.get(tag) if hasattr(exif, 'get') else None
            if v:
                candidates.append(v)
        for tag in tag_numbers:
            v = exif.get(tag) if hasattr(exif, 'get') else None
            if v and v not in candidates:
                candidates.append(v)
        for val in candidates:
            try:
                return datetime.datetime.strptime(str(val), "%Y:%m:%d %H:%M:%S")
            except Exception:
                pass
        return None

    # ── Crea layer ────────────────────────────────────────────────────────────
    from qgis.core import (QgsVectorLayer, QgsFeature,
                            QgsGeometry, QgsPointXY, QgsProject,
                            QgsCoordinateReferenceSystem)

    crs = QgsCoordinateReferenceSystem(crs_epsg)
    layer = QgsVectorLayer(f"Point?crs={crs_epsg}", nome_layer, "memory")
    provider = layer.dataProvider()
    provider.addAttributes([
        crea_campo("name",      "string"),
        crea_campo("Path",      "string"),
        crea_campo("PathUrl",   "string"),
        crea_campo("DateTime",  "string"),
        crea_campo("Title",     "string"),
        crea_campo("Comment",   "string"),
        crea_campo("Note",      "string"),
        crea_campo("Carr",      "string"),
        crea_campo("svincolo",  "string"),
        crea_campo("prog",      "double"),
        crea_campo("descr_1",   "string"),
        crea_campo("foto_num",  "int"),
        crea_campo("Altitude",  "double"),
        crea_campo("Direction", "double"),
    ])
    layer.updateFields()

    totale    = len(file_list)
    inserite  = 0
    saltate   = 0
    no_gps    = 0

    features = []
    for i, path in enumerate(file_list):
        if progress_fn and i % 5 == 0:
            progress_fn(i + 1, totale)

        try:
            img  = PILImage.open(path)
            exif = _get_exif(img)
            img.close()
        except Exception as e:
            log_fn(f"  ⚠ Impossibile leggere {os.path.basename(path)}: {e}")
            saltate += 1
            continue

        # Filtro data
        if data_da or data_a:
            dt = _parse_datetime(exif)
            if dt:
                if data_da and dt.date() < data_da:
                    saltate += 1
                    continue
                if data_a and dt.date() > data_a:
                    saltate += 1
                    continue

        # GPS
        gps = _get_gps(exif)
        if not gps:
            log_fn(f"  ↷ {os.path.basename(path)}: nessun GPS, saltata.")
            no_gps += 1
            continue

        lat = _dms_to_dd(gps.get("GPSLatitude", (0, 0, 0)),
                         gps.get("GPSLatitudeRef", "N"))
        lon = _dms_to_dd(gps.get("GPSLongitude", (0, 0, 0)),
                         gps.get("GPSLongitudeRef", "E"))
        if lat is None or lon is None:
            log_fn(f"  ⚠ {os.path.basename(path)}: coordinate GPS non valide.")
            no_gps += 1
            continue

        alt = None
        try:
            alt_raw = gps.get("GPSAltitude")
            if alt_raw is not None:
                alt = float(alt_raw)
                if gps.get("GPSAltitudeRef") == b'\x01':
                    alt = -alt
        except Exception:
            pass

        direction = None
        try:
            dir_raw = gps.get("GPSImgDirection")
            if dir_raw is not None:
                direction = float(dir_raw)
        except Exception:
            pass

        dt_str = ""
        dt = _parse_datetime(exif)
        if dt:
            dt_str = dt.strftime("%Y-%m-%d %H:%M:%S")

        feat = QgsFeature()
        feat.setGeometry(QgsGeometry.fromPointXY(QgsPointXY(lon, lat)))
        path_url = pathlib.Path(path).as_uri()
        feat.setAttributes([
            os.path.basename(path), path, path_url, dt_str,
            "", "", "", "", "", None, "", None,
            alt, direction
        ])
        features.append(feat)
        inserite += 1

    if progress_fn:
        progress_fn(totale, totale)

    if not features:
        return False, f"Nessuna foto con GPS valido trovata. {no_gps} senza GPS, {saltate} saltate."

    provider.addFeatures(features)
    layer.updateExtents()

    # Imposta Map Tip con anteprima foto
    # Usa replace() per convertire backslash Windows in slash per il tag img
    layer.setMapTipTemplate(
        '<div style="background:white;padding:6px;border:1px solid #aaa;border-radius:6px;max-width:320px;">'
        '<b>[% "name" %]</b><br/>'
        '[% \'<img src="file:///\' || replace("Path", \'\\\\\', \'/\') || \'" style="max-width:300px;max-height:225px;"/>\' %]'
        '<br/><small>[% "DateTime" %]</small>'
        '</div>'
    )
    # Abilita Map Tips automaticamente sul layer
    layer.setCustomProperty("showFeatureCount", False)

    QgsProject.instance().addMapLayer(layer)
    log_fn("  ℹ️  Map Tips attivi: attiva 'Mostra suggerimenti mappa' nella barra degli strumenti per vedere le foto al passaggio del mouse.")

    msg = f"Layer '{nome_layer}' creato con {inserite} foto."
    if no_gps:
        msg += f" {no_gps} senza GPS."
    if saltate:
        msg += f" {saltate} saltate (fuori data o errore lettura)."
    log_fn(msg)
    return True, msg



# ---------------------------------------------------------------------------
# GENERA ETTOMETRICHE
# ---------------------------------------------------------------------------

def genera_ettometriche(nome_layer_punti, campo_km, nome_layer_tratta,
                         passo, tolleranza, nome_output, log_fn, progress_fn=None,
                         max_dist_da_tratta=50.0):

    from qgis.core import (
        QgsVectorLayer, QgsFeature, QgsGeometry, QgsPointXY,
        QgsTextFormat, QgsPalLayerSettings,
        QgsRuleBasedLabeling, QgsTextBufferSettings
    )
    from qgis.PyQt.QtGui import QFont, QColor

    def parse_to_meters(value):
        if value is None or str(value).strip() == "":
            raise ValueError("Valore nullo o vuoto")
        if isinstance(value, (int, float)):
            return int(float(value) * 1000)
        value_str = str(value).replace(",", ".").strip()
        if "+" in value_str:
            parts = value_str.split("+")
            if len(parts) != 2:
                raise ValueError(f"Formato chilometrico non valido: '{value}'")
            try:
                return int(parts[0]) * 1000 + int(parts[1])
            except ValueError:
                raise ValueError(f"Impossibile interpretare il valore chilometrico: '{value}'")
        try:
            return int(float(value_str) * 1000)
        except ValueError:
            raise ValueError(f"Impossibile interpretare il valore chilometrico: '{value}'")

    # Recupera i layer
    layer_punti_list = QgsProject.instance().mapLayersByName(nome_layer_punti)
    if not layer_punti_list:
        return False, f"Layer '{nome_layer_punti}' non trovato."
    layer_punti = layer_punti_list[0]

    layer_tratta_list = QgsProject.instance().mapLayersByName(nome_layer_tratta)
    if not layer_tratta_list:
        return False, f"Layer '{nome_layer_tratta}' non trovato."
    layer_tratta = layer_tratta_list[0]

    if campo_km not in [field.name() for field in layer_punti.fields()]:
        return False, f"Campo '{campo_km}' non trovato nel layer '{nome_layer_punti}'."

    # ── Scegli CRS metrico e converti automaticamente ────────────────────────
    lon_media = _lon_media_layer(layer_punti)
    crs_metrico_authid = _scegli_crs_metrico(lon_media)
    log_fn(f"  CRS di lavoro: {crs_metrico_authid} (longitudine media: {lon_media:.2f}°)")

    t_punti  = _transform_a_metrico(layer_punti.crs(),  crs_metrico_authid)
    t_tratta = _transform_a_metrico(layer_tratta.crs(), crs_metrico_authid)

    if t_punti:
        log_fn(f"  ℹ️  Layer punti  ({layer_punti.crs().authid()}) → {crs_metrico_authid}")
    if t_tratta:
        log_fn(f"  ℹ️  Layer tratta ({layer_tratta.crs().authid()}) → {crs_metrico_authid}")
    if not t_punti and not t_tratta:
        log_fn(f"  ✔ Entrambi i layer già in {crs_metrico_authid}")

    # Geometria della tratta in CRS metrico
    geometrie_tratta = [_geom_metrica(feat.geometry(), t_tratta) for feat in layer_tratta.getFeatures()]
    if len(geometrie_tratta) != 1:
        return False, "Il layer della tratta deve contenere una sola geometria."
    tratta_geom = geometrie_tratta[0]

    punti_esistenti = []
    punti_scartati  = []
    punti_scartati_dist = []  # punti scartati perché troppo lontani dalla tratta

    # Recupera punti esistenti in CRS metrico
    features_list = list(layer_punti.getFeatures())
    totale = len(features_list)
    for i, feat in enumerate(features_list):
        if progress_fn and i % 5 == 0:
            progress_fn(i + 1, totale)
        try:
            raw_value = feat[campo_km]
            if raw_value is None or str(raw_value).strip() == "":
                log_fn(f"  ⚠ Il punto con ID {feat.id()} è stato saltato perché il valore nel campo '{campo_km}' è vuoto.")
                continue
            distanza_m = parse_to_meters(raw_value)
            pt_geom = _geom_metrica(feat.geometry(), t_punti)

            # ── Filtro spaziale: scarta punti troppo lontani dalla polilinea ──
            # Senza questo filtro, lineLocatePoint proietta anche punti molto
            # distanti sugli estremi della tratta, falsando interpolazione ed
            # estrapolazione (causa: range km enorme generato su tratta corta).
            dist_dalla_tratta = pt_geom.distance(tratta_geom)
            if dist_dalla_tratta > max_dist_da_tratta:
                punti_scartati_dist.append((feat.id(), raw_value, dist_dalla_tratta))
                continue

            pt = QgsPointXY(pt_geom.asPoint())
            distanza = tratta_geom.lineLocatePoint(QgsGeometry.fromPointXY(pt))
            pt_proiettato = tratta_geom.interpolate(distanza).asPoint()
            punti_esistenti.append((distanza_m, pt_proiettato, distanza))
        except Exception:
            punti_scartati.append(feat.id())

    # Log dei punti scartati per distanza dalla tratta
    if punti_scartati_dist:
        log_fn(
            f"  ⚠ {len(punti_scartati_dist)} punti scartati perché distano "
            f"più di {max_dist_da_tratta:.0f} m dalla tratta:"
        )
        for fid, val, d in punti_scartati_dist[:10]:
            log_fn(f"     - fid {fid} (km={val}): distanza {d:.1f} m")
        if len(punti_scartati_dist) > 10:
            log_fn(f"     ... e altri {len(punti_scartati_dist) - 10} punti")

    # Ordina punti
    punti_esistenti.sort(key=lambda x: x[0])

    if len(punti_esistenti) < 2:
        msg = (
            f"Servono almeno 2 punti di riferimento validi sulla tratta (trovati {len(punti_esistenti)})."
        )
        if punti_scartati_dist:
            msg += (
                f" {len(punti_scartati_dist)} punti sono stati scartati perché "
                f"distavano più di {max_dist_da_tratta:.0f} m dalla tratta. "
                f"Aumenta il parametro 'Distanza max punti da tratta' se necessario."
            )
        return False, msg

    # ── Rileva la direzione della polilinea rispetto alle progressive ────────
    # Se i km crescono ma le distanze lungo la polilinea decrescono, la tratta
    # è stata disegnata in direzione opposta alle progressive. In tal caso le
    # formule km_inizio = km_primo - dist_primo producono valori assurdi.
    # Soluzione: invertire la geometria della polilinea e specchiare le dist.
    n_cresc = sum(1 for i in range(len(punti_esistenti) - 1)
                  if punti_esistenti[i+1][2] > punti_esistenti[i][2])
    n_decr = sum(1 for i in range(len(punti_esistenti) - 1)
                 if punti_esistenti[i+1][2] < punti_esistenti[i][2])

    if n_decr > n_cresc:
        log_fn(
            "  ↺ La polilinea è orientata in direzione opposta alle progressive "
            "crescenti. Inverto automaticamente la geometria della tratta."
        )
        # Inverti la polilinea (gestisce sia LineString che MultiLineString)
        if tratta_geom.isMultipart():
            parts = tratta_geom.asMultiPolyline()
            reversed_parts = [list(reversed(p)) for p in reversed(parts)]
            tratta_geom = QgsGeometry.fromMultiPolylineXY(reversed_parts)
        else:
            pts = tratta_geom.asPolyline()
            tratta_geom = QgsGeometry.fromPolylineXY(list(reversed(pts)))
        # Le distanze sulla polilinea invertita sono speculari
        L = tratta_geom.length()
        punti_esistenti = [(km, pt, L - dist) for (km, pt, dist) in punti_esistenti]
    elif n_cresc == 0 and n_decr == 0:
        # tutti i punti hanno la stessa dist: caso degenere, ignora
        pass
    elif n_cresc > 0 and n_decr > 0:
        log_fn(
            "  ⚠ I punti di riferimento non sono allineati monotonicamente lungo "
            "la polilinea. Verifica che siano tutti sulla stessa tratta e nella "
            "sequenza corretta."
        )

    # Controllo discrepanze con tolleranza
    for i in range(1, len(punti_esistenti)):
        start_m, _, dist_start = punti_esistenti[i - 1]
        end_m,   _, dist_end   = punti_esistenti[i]

        distanza_calcolata = abs(dist_end - dist_start)
        diff_valore_km     = abs(end_m - start_m)

        start_km = int(start_m // 1000)
        start_metri = round(start_m % 1000)
        end_km   = int(end_m   // 1000)
        end_metri   = round(end_m   % 1000)

        if abs(distanza_calcolata - diff_valore_km) > tolleranza:
            log_fn(
                f"  ⚠ Discrepanza tra la lunghezza del tratto ({distanza_calcolata:.0f} m) "
                f"e la differenza tra le progressive ({diff_valore_km} m) "
                f"tra i km {start_km}+{start_metri:03d} e {end_km}+{end_metri:03d}."
            )

    import math

    # Lunghezza totale della tratta in metrico
    lunghezza_tratta = tratta_geom.length()

    # ── Calcola il km logico all'inizio e alla fine della tratta ─────────────
    # Usando il primo e l'ultimo punto di riferimento come ancora
    dist_primo       = punti_esistenti[0][2]   # distanza lungo tratta del primo punto
    km_primo         = punti_esistenti[0][0]   # km logico del primo punto (metri)
    km_inizio_tratta = km_primo - dist_primo   # km logico all'inizio della tratta (può essere <0)

    dist_ultimo      = punti_esistenti[-1][2]  # distanza lungo tratta dell'ultimo punto
    km_ultimo        = punti_esistenti[-1][0]  # km logico dell'ultimo punto
    km_fine_tratta   = km_ultimo + (lunghezza_tratta - dist_ultimo)

    log_fn(f"  Tratta: km {km_inizio_tratta/1000:.3f} → {km_fine_tratta/1000:.3f} | lunghezza geometrica: {lunghezza_tratta:.0f} m")

    # ── Genera tutti i multipli di passo nel range della tratta ──────────────
    # Primo multiplo >= km_inizio_tratta
    primo_multiplo = math.ceil(km_inizio_tratta / passo) * passo
    # Ultimo multiplo <= km_fine_tratta
    ultimo_multiplo = math.floor(km_fine_tratta / passo) * passo

    km_multipli = []
    km_m = primo_multiplo
    while km_m <= ultimo_multiplo + 0.001:  # 0.001 m di tolleranza floating point
        km_multipli.append(round(km_m))
        km_m += passo

    log_fn(f"  Multipli di {passo} m da inserire: {len(km_multipli)} (da {primo_multiplo:.0f} a {ultimo_multiplo:.0f})")

    # ── Costruisce tabella di interpolazione: km_logico → dist_su_tratta ─────
    # Usa i punti di riferimento come nodi; fuori dai nodi si estrapola linearmente
    def km_a_dist(km_m):
        """Converte un km logico nella distanza corrispondente lungo la tratta."""
        # Prefisso (prima del primo punto di riferimento)
        if km_m <= punti_esistenti[0][0]:
            return dist_primo - (km_primo - km_m)
        # Suffisso (dopo l'ultimo punto di riferimento)
        if km_m >= punti_esistenti[-1][0]:
            return dist_ultimo + (km_m - km_ultimo)
        # Interpolazione tra i punti di riferimento
        for i in range(len(punti_esistenti) - 1):
            km_a, _, dist_a = punti_esistenti[i]
            km_b, _, dist_b = punti_esistenti[i + 1]
            if km_a <= km_m <= km_b:
                frac = (km_m - km_a) / (km_b - km_a)
                return dist_a + frac * (dist_b - dist_a)
        return None

    # ── Costruisce la lista punti_completi con solo multipli di passo ─────────
    punti_completi = []
    for km_m in km_multipli:
        dist = km_a_dist(km_m)
        if dist is None:
            continue
        # Clamp alla lunghezza della tratta (tolleranza numerica)
        dist = max(0.0, min(dist, lunghezza_tratta))
        punti_completi.append((km_m, dist))


    log_fn(f"  {len(punti_completi)} punti generati.")

    # Set dei km logici di riferimento (punti forniti dall'utente) per determinare tipo
    km_riferimento = {pe[0] for pe in punti_esistenti}
    km_primo_rif   = punti_esistenti[0][0]
    km_ultimo_rif  = punti_esistenti[-1][0]

    # Layer di output in CRS metrico
    out_layer = QgsVectorLayer("Point?crs=" + crs_metrico_authid, nome_output, "memory")
    dp = out_layer.dataProvider()
    dp.addAttributes([
        crea_campo("ID",             "int"),
        crea_campo("km_intero",      "int"),
        crea_campo("prog",           "string"),
        crea_campo("prog_metri",     "int"),
        crea_campo("dist_su_tratta", "double"),
        crea_campo("tipo",           "string"),
    ])
    out_layer.updateFields()

    # Aggiungi punti al layer
    id_counter = 1
    feats = []
    for km_val_m, dist in punti_completi:
        if progress_fn:
            progress_fn(id_counter, len(punti_completi))
        pt = QgsPointXY(tratta_geom.interpolate(dist).asPoint())
        km_int  = int(km_val_m // 1000)
        metri   = round(km_val_m % 1000)
        prog_label = f"{km_int}+{metri:03d}"

        # Determina tipo
        if km_val_m in km_riferimento:
            tipo = "esistente"
        elif km_val_m < km_primo_rif or km_val_m > km_ultimo_rif:
            tipo = "estrapolata"
        else:
            tipo = "interpolata"

        f = QgsFeature()
        f.setGeometry(QgsGeometry.fromPointXY(pt))
        f.setAttributes([id_counter, km_int, prog_label, int(km_val_m), round(dist, 2), tipo])
        feats.append(f)
        id_counter += 1

    dp.addFeatures(feats)
    out_layer.updateExtents()
    QgsProject.instance().addMapLayer(out_layer)

    # Etichettatura rule-based
    format_km = QgsTextFormat()
    format_km.setFont(QFont("Arial", 11))
    format_km.setSize(11)
    format_km.setColor(QColor('#f3072e'))
    buffer_km = QgsTextBufferSettings()
    buffer_km.setEnabled(True)
    buffer_km.setSize(1.0)
    buffer_km.setColor(QColor('#f3072e'))
    format_km.setBuffer(buffer_km)

    settings_km = QgsPalLayerSettings()
    settings_km.fieldName = "prog"
    settings_km.format = format_km
    settings_km.placement = placement_around_point()
    settings_km.enabled = True
    rule_km = QgsRuleBasedLabeling.Rule(settings_km)
    rule_km.setFilterExpression("right(\"prog\", 3) = '000'")

    format_interp = QgsTextFormat()
    format_interp.setFont(QFont("Arial", 8))
    format_interp.setSize(8)
    format_interp.setColor(QColor("#f3072e"))

    settings_interp = QgsPalLayerSettings()
    settings_interp.fieldName = "prog"
    settings_interp.format = format_interp
    settings_interp.placement = placement_around_point()
    settings_interp.enabled = True
    rule_interp = QgsRuleBasedLabeling.Rule(settings_interp)
    rule_interp.setFilterExpression("right(\"prog\", 3) != '000'")

    root_rule = QgsRuleBasedLabeling.Rule(None)
    root_rule.appendChild(rule_km)
    root_rule.appendChild(rule_interp)

    labeling = QgsRuleBasedLabeling(root_rule)
    out_layer.setLabeling(labeling)
    out_layer.setLabelsEnabled(True)
    out_layer.setCustomProperty("labeling", "rule-based")
    out_layer.triggerRepaint()

    msg = f"Ettometriche generate: layer '{nome_output}' con {len(punti_completi)} punti."
    if punti_scartati:
        msg += f" {len(punti_scartati)} punti scartati: {', '.join(map(str, punti_scartati))}."
    log_fn(f"  ✔ {msg}")
    return True, msg


# ---------------------------------------------------------------------------
# CALCOLO AUTOMATICO CARREGGIATA
# ---------------------------------------------------------------------------

def calcola_carreggiata_automatica(nome_layer_foto, nome_layer_tratta_a, nome_layer_tratta_b,
                                    val_carr_a, val_carr_b, log_fn):
    """
    Per ogni foto calcola la distanza minima dalla polilinea A e B nel CRS delle foto.
    - Salta se la distanza minima > 50 m (foto fuori tratta)
    - Salta se il campo svincolo è compilato con un nome reale
    - Scrive il risultato in Carr_auto
    - Confronta Carr_auto con Carr (case-insensitive) e restituisce warning se diversi
    """
    from qgis.core import (QgsCoordinateTransform, QgsCoordinateTransformContext,
                            QgsGeometry)

    SOGLIA_DISTANZA = 50  # metri

    # ── Carica layer ─────────────────────────────────────────────────────────
    foto_layers   = QgsProject.instance().mapLayersByName(nome_layer_foto)
    tratta_a_list = QgsProject.instance().mapLayersByName(nome_layer_tratta_a)
    tratta_b_list = QgsProject.instance().mapLayersByName(nome_layer_tratta_b)

    if not foto_layers:
        return False, f"Layer foto '{nome_layer_foto}' non trovato."
    if not tratta_a_list:
        return False, f"Layer tratta A '{nome_layer_tratta_a}' non trovato."
    if not tratta_b_list:
        return False, f"Layer tratta B '{nome_layer_tratta_b}' non trovato."

    foto_layer     = foto_layers[0]
    layer_tratta_a = tratta_a_list[0]
    layer_tratta_b = tratta_b_list[0]

    crs_foto = foto_layer.crs()

    # ── Scegli CRS metrico e riproietta le tratte ─────────────────────────────
    lon_media = _lon_media_layer(foto_layer)
    crs_metrico_authid = _scegli_crs_metrico(lon_media)
    log_fn(f"  ℹ️  CRS di lavoro: {crs_metrico_authid} (longitudine media: {lon_media:.2f}°)")

    t_foto = _transform_a_metrico(crs_foto, crs_metrico_authid)

    def _geom_layer_in_metrico(layer):
        t = _transform_a_metrico(layer.crs(), crs_metrico_authid)
        if t:
            log_fn(f"  ℹ️  Tratta '{layer.name()}': {layer.crs().authid()} → {crs_metrico_authid}")
        geoms = []
        for feat in layer.getFeatures():
            g = _geom_metrica(feat.geometry(), t)
            geoms.append(g)
        return QgsGeometry.unaryUnion(geoms)

    geom_a = _geom_layer_in_metrico(layer_tratta_a)
    geom_b = _geom_layer_in_metrico(layer_tratta_b)

    if geom_a is None or geom_a.isEmpty():
        return False, f"Layer tratta A '{nome_layer_tratta_a}' non contiene geometrie valide."
    if geom_b is None or geom_b.isEmpty():
        return False, f"Layer tratta B '{nome_layer_tratta_b}' non contiene geometrie valide."

    log_fn(f"  ℹ️  CRS di lavoro: {crs_metrico_authid} | Soglia distanza: {SOGLIA_DISTANZA} m")

    # ── Aggiungi campo Carr_auto se mancante ──────────────────────────────────
    field_names = [f.name() for f in foto_layer.fields()]
    if "Carr_auto" not in field_names:
        foto_layer.dataProvider().addAttributes([crea_campo("Carr_auto", "string")])
        foto_layer.updateFields()
        field_names = [f.name() for f in foto_layer.fields()]

    idx_carr_auto = foto_layer.fields().indexFromName("Carr_auto")
    idx_svincolo  = foto_layer.fields().indexFromName("svincolo") if "svincolo" in field_names else -1
    idx_carr_orig = foto_layer.fields().indexFromName("Carr")     if "Carr"     in field_names else -1

    # ── Calcola distanza e assegna ────────────────────────────────────────────
    foto_layer.startEditing()
    assegnate_a = assegnate_b = saltate_dist = saltate_sv = errori = 0
    features = list(foto_layer.getFeatures())
    discrepanze = []  # (fid, nome_file, carr_orig, carr_auto)

    for foto in features:
        geom = foto.geometry()
        if geom is None or geom.isEmpty():
            errori += 1
            continue

        # Salta se svincolo compilato
        if idx_svincolo >= 0:
            sv = str(foto["svincolo"] or "").strip()
            if sv and sv.lower() not in ("", "no", "null", "none", "sospetto"):
                log_fn(f"  ↷ fid {foto.id()}: svincolo '{sv}', carreggiata non calcolata.")
                saltate_sv += 1
                continue

        # Calcola distanze in CRS metrico
        geom_calc = _geom_metrica(geom, t_foto)
        dist_a = geom_calc.distance(geom_a)
        dist_b = geom_calc.distance(geom_b)
        dist_min = min(dist_a, dist_b)

        # Salta se troppo lontano da entrambe le tratte
        if dist_min > SOGLIA_DISTANZA:
            log_fn(f"  ↷ fid {foto.id()}: distanza minima {dist_min:.1f} m > {SOGLIA_DISTANZA} m, saltata.")
            saltate_dist += 1
            continue

        carr_auto = val_carr_a if dist_a <= dist_b else val_carr_b
        foto_layer.changeAttributeValue(foto.id(), idx_carr_auto, carr_auto)

        if dist_a <= dist_b:
            assegnate_a += 1
        else:
            assegnate_b += 1

        # Confronto con Carr originale (solo se compilato)
        if idx_carr_orig >= 0:
            carr_orig = str(foto["Carr"] or "").strip()
            if carr_orig and carr_orig.lower() not in ("", "null", "none"):
                if carr_orig.lower() != carr_auto.lower():
                    nome_file = str(foto["name"] or foto.id())
                    discrepanze.append((foto.id(), nome_file, carr_orig, carr_auto))

    foto_layer.commitChanges()

    # ── Aggiorna il layer in QGIS ─────────────────────────────────────────────
    foto_layer.triggerRepaint()

    # ── Log riassuntivo ───────────────────────────────────────────────────────
    msg = (f"Carr_auto calcolata: {assegnate_a} → '{val_carr_a}', {assegnate_b} → '{val_carr_b}'.")
    if saltate_sv:
        msg += f" {saltate_sv} saltate (svincolo)."
    if saltate_dist:
        msg += f" {saltate_dist} saltate (distanza > {SOGLIA_DISTANZA} m)."
    if errori:
        msg += f" {errori} errori (senza geometria)."
    log_fn(f"  ✔ {msg}")

    # ── Warning discrepanze ───────────────────────────────────────────────────
    if discrepanze:
        log_fn(f"\n  ⚠ ATTENZIONE — {len(discrepanze)} foto con Carr diverso da Carr_auto:")
        for fid, nome, orig, auto in discrepanze:
            log_fn(f"    • fid {fid} | {nome} | Carr='{orig}' → Carr_auto='{auto}'")
        warning_msg = (
            f"⚠ {len(discrepanze)} foto hanno il campo Carr diverso da Carr_auto:\n\n" +
            "\n".join(f"• {nome}  (Carr='{orig}' → Carr_auto='{auto}')"
                      for _, nome, orig, auto in discrepanze)
        )
        return True, msg + "\n\nWARNING:" + warning_msg

    return True, msg
